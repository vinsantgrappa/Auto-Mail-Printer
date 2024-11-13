import csv
import html
import imaplib
import email
from email import policy
from email.parser import BytesParser
import win32com.client
import random
import base64
import os
import shutil
import PyPDF2
import glob
import sys
import logging
from PIL import Image
from PyPDF2.errors import FileNotDecryptedError
from docx import Document
import datetime
import io
import time
import traceback
import re

class AutoMailPrinter:
    """
    メールを自動で取得し、添付ファイルや本文を処理してPDFに変換し、
    所定のフォルダに保存するクラス。
    """

    def __init__(self):
        # ロガーの設定
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.DEBUG)

        # ログファイルのハンドラーを生成
        log_file_path = os.path.join("path_to_logs", "autoprint_log.txt")
        self.h.setLevel(logging.DEBUG)

        # ログのフォーマッタを生成
        self.fmt = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')

        # ハンドラーにフォーマッタを設定
        self.h.setFormatter(self.fmt)

        # ロガーにハンドラーを設定
        self.logger.addHandler(self.h)

        # ファイルパス
        self.log_data = os.path.join("path_to_your_logs", "autoprint_log.txt")
        self.after_moved_error = os.path.join("path_to_fax_folder", "auto_print_stop_notice.txt")
        self.error_text = os.path.join("path_to_error_logs", "auto_print_error_notice.txt")
        self.customer_address = os.path.join("path_to_data_files", "customer_address.csv")
        self.after_atmark = os.path.join("path_to_data_files", "after_atmark.csv")
        self.mail_body_to_txt_file = os.path.join("path_to_output_files", "mail_body_to_txt_file.txt")
        self.mail_body_to_word_file = os.path.join("path_to_output_files", "mail_body_to_word_file.docx")
        self.mail_body_to_pdf_file = os.path.join("path_to_pdf_files", "01-test.pdf")
        self.PDF_dir = os.path.join("path_to_pdf_files")
        self.received_folder_dir = os.path.join("path_to_received_fax")

        # メールのエンコードリスト
        self.encode_list = ["UTF-8", "iso-2022-jp", "ISO-2022-JP"]

        # メールアドレスの前処理用の正規表現パターン
        self.erase_pattern_email_add = r'^.*?"'

        # Gmailのログイン情報（環境変数から取得）
        self.UserName = os.environ["gmail_address_mail_print"]
        self.PassName = os.environ["gmail_password_mail_print"]

        # カウンター変数の初期化
        self.n = 0

        # フラグ変数の初期化
        self.delete_flag_excel = False
        self.delete_flag_zip = False
        self.delete_flag_png = False
        self.delete_flag_dat = False
        self.delete_flag = False
        self.delete_tif_flag = False
        self.pdf_flg = False
        self.company_name = None
        self.mail_add_found = False
        self.body_list = []

    def exit(self):
        """
        プログラムの終了処理を行うメソッド。
        """
        self.logger.info("処理が終了")
        self.logger.info(traceback.format_exc())

    def get_pdf(self, d):
        """
        添付ファイルからPDFを取得し、必要に応じて変換・削除を行うメソッド。

        Parameters:
            d (list): メールデータ。
        """
        try:
            # メールメッセージを解析
            email_message = email.message_from_bytes(d[0][1])
            for part in email_message.walk():
                content_type = part.get_content_type()
                file_name = part.get_filename()

                if not file_name:
                    continue

                # ファイル名のデコード処理
                fns = file_name.split("?")
                try:
                    output_file_name = base64.b64decode(fns[-2]).decode(fns[1])
                except Exception:
                    output_file_name = file_name

                print("content_type：", content_type)

                # ファイルの拡張子に応じてファイル名を設定
                if "application/pdf" in content_type:
                    output_file_name = "attached_file" + str(random.randrange(1000)) + ".pdf"
                if "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" in content_type:
                    output_file_name = "attached_file" + str(random.randrange(1000)) + ".xlsx"
                if "application/x-zip-compressed" in content_type:
                    output_file_name = "attached_file" + str(random.randrange(1000)) + ".zip"

                # 添付ファイルを保存
                with open(f'{os.getcwd()}/{output_file_name}', 'wb') as f:
                    try:
                        f.write(part.get_payload(decode=True))
                    except TypeError:
                        continue
                    f.close()

                    # ファイルの拡張子を取得
                    extention = output_file_name[-3:]
                    print(f"拡張子：{extention}")
                    print(f"添付ファイル：{output_file_name}")
                    pdf_path = os.path.join(os.getcwd(), output_file_name)
                    after_moved = os.path.join(self.PDF_dir, "02-" + output_file_name)
                    convert_pdf = os.path.join(self.PDF_dir, "02-" + str(random.randrange(1000)) + ".pdf")

                    # ファイルを指定のディレクトリに移動
                    shutil.copy(pdf_path, after_moved)
                    os.remove(pdf_path)

                    # 拡張子を小文字で取得
                    ext = output_file_name.lower().split('.')[-1]

                    # 拡張子に応じて処理を分岐
                    if ext == "xlsx":
                        os.remove(after_moved)
                        self.delete_flag_excel = True

                    elif ext == "zip":
                        os.remove(after_moved)
                        self.delete_flag_zip = True

                    elif ext in ['png', 'jpg', 'jpeg']:
                        # 画像をPDFに変換
                        self.convert_img_to_pdf(after_moved, convert_pdf)

                    elif ext in ['tif', 'tiff']:
                        os.remove(after_moved)
                        self.delete_tif_flag = True

                    elif ext == 'dat':
                        os.remove(after_moved)
                        self.delete_flag_dat = True

                    elif ext != 'pdf':
                        os.remove(after_moved)
                        self.delete_flag = True

        except Exception as e:
            self.error_log(e)

    def convert_img_to_pdf(self, after_moved, convert_pdf):
        """
        画像ファイルをPDFに変換するメソッド。

        Parameters:
            after_moved (str): 変換前のファイルパス。
            convert_pdf (str): 変換後のPDFファイルパス。
        """
        image1 = Image.open(after_moved)
        img_size = image1.size

        # 画像サイズが大きい場合はリサイズ
        if img_size > (3000, 3000):
            image1 = image1.resize((1280, 720))
            image1.save(after_moved)

        im_pdf = image1.convert("RGB")
        im_pdf.save(convert_pdf)
        image1.close()
        print(f"画像⇒pdfへ変換：{convert_pdf}")
        os.remove(after_moved)
        self.delete_flag_png = True

    def error_log(self, e):
        """
        エラーログを記録し、プログラムを終了するメソッド。

        Parameters:
            e (Exception): 発生した例外。
        """
        print(e)
        self.logger.error(sys.exc_info())
        print(sys.exc_info())
        with open(self.error_text, "w") as f:
            f.write(str(self.logger.error(sys.exc_info())))
            f.close()
        shutil.copy(self.log_data, self.after_moved_error)
        sys.exit()

    def compare_body(self, body_sentence):
        """
        メール本文が既に処理済みかをチェックするメソッド。

        Parameters:
            body_sentence (list): メール本文をスペースで分割したリスト。

        Returns:
            int or None: メールアドレスのインデックス、またはNone。
        """
        try:
            if not self.body_list:
                self.body_list.append(body_sentence)
                mail_address = body_sentence.index("To:") - 1
                return mail_address
            else:
                if body_sentence not in self.body_list:
                    self.body_list.append(body_sentence)
                    print("新規で本文をリストに追加しました。")
                    mail_address = body_sentence.index("To:") - 1
                    return mail_address
                else:
                    print("ダブりが見つかりました！")
                    mail_address = None
                    return mail_address
        except Exception as e:
            print(e)

    def convert_body_to_text(self, email_add, body2):
        """
        メール本文をテキストファイルに書き込むメソッド。

        Parameters:
            email_add (str): メールアドレス。
            body2 (str): メール本文。

        Returns:
            str: テキストファイルの内容。
        """
        with open(self.mail_body_to_txt_file, "w", encoding="utf-8", newline='') as f:
            f.write(email_add)
            f.write("\n\n")
            for row in body2:
                f.write(row)
            f.close()

        with open(self.mail_body_to_txt_file, "r", encoding="utf-8") as f:
            text = f.read()
            f.close()
            return text

    def identify_company_name(self, fixed_mail_address):
        """
        メールアドレスから会社名を特定するメソッド。

        Parameters:
            fixed_mail_address (str): 整形済みのメールアドレス。

        Returns:
            str: 会社名。
        """
        # 顧客アドレスリストから検索
        with open(self.customer_address, "r", encoding="shift-jis") as f:
            for row in csv.DictReader(f):
                if fixed_mail_address in row["E-mail"]:
                    self.company_name = str(row["会社名"])
                    print("----", self.company_name, "----")
                    self.mail_add_found = True
                    return self.company_name

        # ドメインのみで検索
        if not self.mail_add_found:
            parts = fixed_mail_address.split("@")
            if len(parts) > 1:
                after_at = "@" + parts[1]
                with open(self.after_atmark, "r", encoding="shift-jis") as f2:
                    for row2 in csv.DictReader(f2):
                        if after_at in row2["E-mail"]:
                            self.company_name = "未登録のアドレス"
                            return self.company_name
                f2.close()
            else:
                print(f"no parts in {fixed_mail_address}")
        f.close()

    def process_text_plain(self, msg_parser, msg_encoding, email_add):
        """
        content_typeが"text/plain"の場合のメール本文処理メソッド。

        Parameters:
            msg_parser (email.message.Message): メールメッセージオブジェクト。
            msg_encoding (str): メッセージのエンコーディング。
            email_add (str): メールアドレス。
        """
        try:
            print(f"content type: {msg_parser.get_content_type()}")
            body = msg_parser.get_payload(decode=True)

            if msg_encoding == 'utf-8':
                try:
                    body = base64.urlsafe_b64decode(body.encode('ASCII')).decode("utf-8")
                except Exception:
                    print(traceback.format_exc())
                    body = msg_parser.get_payload()
            else:
                print(f"{msg_encoding}にてpayload")
                body = body.decode(f'{msg_encoding}', errors="ignore")

            # HTMLエンティティのデコードと不要な文字の削除
            body = html.unescape(body)
            body = re.sub(r'[\x00-\x09\x0b-\x0c\x0e-\x1f\x7f]', '', body)
            body2 = body.replace("\xa0", "")
            body2 = body2.replace("=", "＝")
            print(f"メール本文{body2}")

            split_body = body2.split()
            mail_address = self.compare_body(split_body)

            try:
                fixed_mail_address = split_body[mail_address].replace("<", "").replace(">", "")
                fixed_mail_address = re.sub(self.erase_pattern_email_add, "", fixed_mail_address)
            except TypeError:
                fixed_mail_address = "ダブりメール"

            print("メール送信者:", fixed_mail_address)

            # 会社名の特定
            self.identify_company_name(fixed_mail_address)

            # メール本文をWordからPDFに変換
            wdFormatPDF = 17
            document = Document()
            document.add_paragraph(self.convert_body_to_text(email_add, body2))
            document.save(self.mail_body_to_word_file)
            input_file = self.mail_body_to_word_file
            output_file = self.mail_body_to_pdf_file

            try:
                word = win32com.client.Dispatch("Word.Application")
            except Exception as e:
                print(e)
                word = win32com.client.Dispatch("Word.Application")

            doc = word.Documents.Open(input_file)
            doc.SaveAs(output_file, FileFormat=wdFormatPDF)
            doc.Close()
            word.Quit()

            # PDFファイルのマージ
            files = glob.glob(self.PDF_dir + "/*")
            merger = PyPDF2.PdfMerger(strict=False)
            for file in files:
                merger.append(file)
            try:
                merger.write(
                    os.path.join(self.received_folder_dir, self.company_name + "【メール依頼】-" + str(
                        random.randrange(1000)) + ".pdf"))
            except Exception:
                print("リストに存在しないアドレス：" + fixed_mail_address)
                self.logger.error("リストに存在しないアドレス：" + fixed_mail_address)
                merger.close()
            merger.close()

            # 一時ファイルの削除
            for file in files:
                os.remove(file)

            print("\nメールサーチ中、、、\n")

        except Exception as e:
            self.error_log(e)

    def process_multipart_alternative(self, msg_parser, msg_encoding, email_add):
        """
        content_typeが"multipart/alternative"の場合のメール本文処理メソッド。

        Parameters:
            msg_parser (email.message.Message): メールメッセージオブジェクト。
            msg_encoding (str): メッセージのエンコーディング。
            email_add (str): メールアドレス。
        """
        try:
            print(f"content type: {msg_parser.get_content_type()}")
            for payload in msg_parser.get_payload():
                if payload.get_content_type() == "text/plain":
                    body = payload.get_payload(decode=True)
                    if msg_encoding == 'utf-8':
                        try:
                            body = base64.urlsafe_b64decode(body.encode('ASCII')).decode("utf-8")
                        except Exception:
                            body = payload.get_payload()
                    else:
                        print(f"{msg_encoding}にてpayload")
                        body = body.decode(f'{msg_encoding}', errors="ignore")
                    body = html.unescape(body)
                    body = re.sub(r'[\x00-\x09\x0b-\x0c\x0e-\x1f\x7f]', '', body)
                    body2 = body.replace("\xa0", "")
                    body2 = body2.replace("=", "＝")
                    split_body = body2.split()
                    print(f"メール本文：{body2}")
                    mail_address = self.compare_body(split_body)

                    try:
                        fixed_mail_address = split_body[mail_address].replace("<", "").replace(">", "")
                        fixed_mail_address = re.sub(self.erase_pattern_email_add, "", fixed_mail_address)
                    except TypeError:
                        fixed_mail_address = "ダブりメール"

                    print("メール送信者:", fixed_mail_address)

                    # 会社名の特定
                    self.identify_company_name(fixed_mail_address)

                    # メール本文をWordからPDFに変換
                    wdFormatPDF = 17
                    document = Document()
                    document.add_paragraph(self.convert_body_to_text(email_add, body2))
                    document.save(self.mail_body_to_word_file)
                    input_file = self.mail_body_to_word_file
                    output_file = self.mail_body_to_pdf_file

                    try:
                        word = win32com.client.Dispatch("Word.Application")
                    except Exception as e:
                        print(f"{e}")
                        word = win32com.client.Dispatch("Word.Application")

                    doc = word.Documents.Open(input_file)
                    doc.SaveAs(output_file, FileFormat=wdFormatPDF)
                    doc.Close()
                    word.Quit()

                    # PDFファイルのマージ
                    files = glob.glob(self.PDF_dir + "/*")
                    merger = PyPDF2.PdfMerger(strict=False)
                    for file in files:
                        print(f"結合対象のファイル：{file}")
                        merger.append(file)
                    try:
                        merger.write(
                            os.path.join(self.received_folder_dir, self.company_name + "【メール依頼】-" + str(
                                random.randrange(1000)) + ".pdf"))
                    except Exception:
                        print("リストに存在しないアドレス：" + fixed_mail_address)
                        self.logger.error("リストに存在しないアドレス：" + fixed_mail_address)
                        merger.close()
                    merger.close()

                    # 一時ファイルの削除
                    for file in files:
                        os.remove(file)

                    print("\nメールサーチ中、、、\n")

        except Exception as e:
            self.error_log(e)

    def process_multipart_mixed(self, d, msg_parser, email_add):
        """
        添付ファイル付きメールの処理メソッド。

        Parameters:
            d (list): メールデータ。
            msg_parser (email.message.Message): メールメッセージオブジェクト。
            email_add (str): メールアドレス。
        """
        try:
            print(f"content type: {msg_parser.get_content_type()}")
            email_message = email.message_from_bytes(d[0][1])

            for part in email_message.walk():
                file_name = part.get_filename()

                if part.get_content_maintype() == "multipart":
                    continue

                if not file_name:
                    # メール本文の処理
                    charset = part.get_content_charset() or "utf-8"
                    body = part.get_payload(decode=True).decode(charset, errors="replace")
                    body = html.unescape(body)
                    body2 = body.replace("\xa0", "")
                    body2 = body2.replace("=", "＝")
                    split_body = body2.split()
                    print(f"メール本文：{body2}")
                    mail_address = self.compare_body(split_body)

                    try:
                        fixed_mail_address = split_body[mail_address].replace("<", "").replace(">", "")
                        fixed_mail_address = re.sub(self.erase_pattern_email_add, "", fixed_mail_address)
                    except TypeError:
                        fixed_mail_address = "ダブりメール"

                    print("メール送信者:", fixed_mail_address)

                    # 会社名の特定
                    self.identify_company_name(fixed_mail_address)

                    # メール本文をWordからPDFに変換
                    wdFormatPDF = 17
                    document = Document()
                    document.add_paragraph(self.convert_body_to_text(email_add, body2))
                    document.save(self.mail_body_to_word_file)
                    input_file = self.mail_body_to_word_file
                    output_file = self.mail_body_to_pdf_file

                    try:
                        word = win32com.client.Dispatch("Word.Application")
                    except Exception as e:
                        print(f"Word's COM Objectへの接続が失敗しました。  {e}")
                        word = win32com.client.Dispatch("Word.Application")

                    doc = word.Documents.Open(input_file)
                    doc.SaveAs(output_file, FileFormat=wdFormatPDF)
                    doc.Close()
                    word.Quit()

                    # PDFファイルのマージ
                    files = glob.glob(self.PDF_dir + "/*")
                    merger = PyPDF2.PdfMerger(strict=False)
                    for file in files:
                        print("添付ファイル名：", file)
                        try:
                            merger.append(file)
                        except FileNotDecryptedError:
                            self.pdf_flg = True

                    # フラグに応じてファイル名を変更して保存
                    if self.delete_flag_excel:
                        output_filename = self.company_name + "【メール依頼 EXCEL削除済み】-" + str(random.randrange(1000)) + ".pdf"
                    elif self.delete_flag_zip:
                        output_filename = self.company_name + "【メール依頼 ZIPファイル削除済み】-" + str(random.randrange(1000)) + ".pdf"
                    elif self.delete_flag_png:
                        output_filename = self.company_name + "【メール依頼 画像⇒PDF変換済み】-" + str(random.randrange(1000)) + ".pdf"
                    elif self.delete_flag_dat:
                        output_filename = self.company_name + "【メール依頼 DATファイル削除済み】-" + str(random.randrange(1000)) + ".pdf"
                    elif self.delete_flag:
                        output_filename = self.company_name + "【メール依頼 添付不可ファイル削除済み】-" + str(random.randrange(1000)) + ".pdf"
                    elif self.pdf_flg:
                        output_filename = self.company_name + "【メール依頼 パスワード保護PDF削除済み】-" + str(random.randrange(1000)) + ".pdf"
                    elif self.delete_tif_flag:
                        output_filename = self.company_name + "【メール依頼 TIFファイル削除済み】-" + str(random.randrange(1000)) + ".pdf"
                    else:
                        output_filename = self.company_name + "【メール依頼】-" + str(random.randrange(1000)) + ".pdf"

                    try:
                        merger.write(os.path.join(self.received_folder_dir, output_filename))
                    except Exception:
                        print("リストに存在しないアドレス：" + fixed_mail_address)
                        self.logger.error("リストに存在しないアドレス：" + fixed_mail_address)
                        merger.close()
                    merger.close()

                    # 一時ファイルの削除
                    for file in files:
                        os.remove(file)
                    print("\n出力完了！\n")
                    print("メールサーチ中、、、\n")
                    break
        except Exception as e:
            self.error_log(e)

    def run(self):
        """
        メイン処理を行うメソッド。
        Gmailにログインし、未読メールをチェックして処理を行う。
        """
        # ログイン処理
        gmail = imaplib.IMAP4_SSL("imap.gmail.com", "993")
        gmail.login(self.UserName, self.PassName)
        print(f"\n{gmail}　ログインが完了しました\n")
        print("メールサーチ中、、、")

        while True:
            now = datetime.datetime.now()
            hour_data = now.strftime("%H")
            min_data = now.strftime("%M")
            calc_time = int(hour_data + min_data)

            try:
                gmail.select()
                if calc_time > 2030:
                    sys.exit()
                else:
                    gmail.select()
                    head, data = gmail.search(None, "UNSEEN")  # 未読メールのみ取得
            except Exception as e:
                print(f"error:{e} が発生しました。リトライを開始します。")
                self.logger.error(sys.exc_info())
                failed_flg = True
                retry_delay = 5
                while failed_flg:

                    try:
                        gmail = imaplib.IMAP4_SSL("imap.gmail.com", "993")
                        gmail.login(self.UserName, self.PassName)
                        print(f"\n{gmail}　ログインが完了しました\n")
                        print("メールサーチ中、、、")
                        gmail.select()
                        head, data = gmail.search(None, "UNSEEN")  # 未読メールのみ取得
                        failed_flg = False
                    except Exception as e:
                        print(f"error:{e}　失敗しました。5秒後にリトライします。")
                        time.sleep(retry_delay)

            self.company_name = None
            self.n += 1

            for num in data[0].split():
                try:
                    h, d = gmail.fetch(num, "(RFC822)")
                except Exception as e:
                    self.error_log(e)

                raw_email = d[0][1]

                # charsetを取得
                code = None
                check_charset = email.message_from_string(raw_email.decode("utf-8"))
                for encoding in self.encode_list:
                    if str(encoding) in str(check_charset):
                        code = encoding
                        print("charset = ", code)

                # 取得したcharsetでメール本文を取得
                stream = io.BytesIO(raw_email)
                msg_parser = BytesParser(policy=policy.default).parse(stream)
                msg = email.message_from_string(raw_email.decode(str(code), errors="replace"))
                msg_encoding = email.header.decode_header(msg.get("Subject"))[0][1] or "iso-2022-jp"

                # タイトルの情報を抽出
                msg_subject = email.header.decode_header(msg.get("Subject"))[0][0]

                # エンコーディング
                try:
                    subject = str(msg_subject.decode(msg_encoding, errors="replace"))
                except Exception:
                    subject = "タイトルなし"
                str_msg = str(msg["From"])
                rep_msg = str_msg.replace("<", "")
                rep_msg = rep_msg.replace(">", "")
                split_str = rep_msg.split()

                try:
                    email_add = split_str[1]
                except Exception as e:
                    self.error_log(e)

                print("転送者：", email_add)
                print("メールタイトル：", subject)

                # content_typeに応じて処理を分岐
                if msg_parser.get_content_type() == "multipart/alternative":
                    self.process_multipart_alternative(msg_parser, msg_encoding, email_add)

                if msg_parser.get_content_type() == "text/plain":
                    self.process_text_plain(msg_parser, msg_encoding, email_add)

                if msg.get_content_type() == "multipart/mixed":
                    self.get_pdf(d)
                    try:
                        self.process_multipart_mixed(d, msg_parser, email_add)
                    except Exception as e:
                        print(e)
                        self.logger.error(sys.exc_info())
                        merger = PyPDF2.PdfMerger(strict=False)
                        merger.close()
                        files = glob.glob(self.PDF_dir + "/*")
                        for file in files:
                            os.remove(file)
                    break

            # フラグのリセット
            self.delete_flag_excel = False
            self.delete_flag_zip = False
            self.delete_flag_png = False
            self.delete_flag_dat = False
            self.delete_flag = False
            self.delete_tif_flag = False
            self.pdf_flg = False
            self.mail_add_found = False

            # 次のチェックまで待機
            t = 15
            while t > 0:
                t -= 1
                time.sleep(1)

if __name__ == '__main__':
    printer = AutoMailPrinter()
    printer.run()
